from flask import Flask, render_template, request, jsonify
import PyPDF2
import io
import os
import docx
from docx import Document
from werkzeug.utils import secure_filename
from zipfile import ZipFile
from xml.dom.minidom import parseString
from urllib.request import urlopen
from flask import Flask, jsonify
from flask_cors import CORS






def extract_text_from_odt(odt_path):
    """Extracts text from an ODT file.

    Args:
        odt_path (str): The path to the ODT file.

    Returns:
        str: A string containing the text of the ODT file.
    """
    with ZipFile(odt_path, 'r') as odt_file:
        with odt_file.open('content.xml') as content_file:
            content = content_file.read()

    dom = parseString(content)
    paragraphs = dom.getElementsByTagName('text:p')

    text = ""
    for paragraph in paragraphs:
        for node in paragraph.childNodes:
            if node.nodeType == node.TEXT_NODE:
                text += node.data

    return text



def extract_text_from_pdf(pdf_file):
  """Extracts text from a PDF file.

  Args:
    pdf_file: The path to the PDF file.

  Returns:
    A string containing the text of the PDF file.
  """
  with open(pdf_file, "rb") as f:
    pdf_reader = PyPDF2.PdfReader(f)
    text = ""
    for page in pdf_reader.pages:
      text += page.extract_text()
    return text

def extract_text_from_txt(txt_file):
  """Extracts text from a TXT file.

  Args:
    txt_file: The path to the TXT file.

  Returns:
    A string containing the text of the TXT file.
  """
  with open(txt_file, "r") as f:
    text = f.read()
    return text

def extract_text_from_doc(doc_file):
  """Extracts text from a DOC file.

  Args:
    doc_file: The path to the DOC file.

  Returns:
    A string containing the text of the DOC file.
  """

  with open(doc_file, "rb") as f:
    doc = docx.Document(f)
    text = ""
    for paragraph in doc.paragraphs:
      text += paragraph.text
    return text



from docx import Document

def extract_text_from_docx(docx_file):
    """Extracts text from a DOCX file.

    Args:
        docx_file (str): The path to the DOCX file.

    Returns:
        str: A string containing the text of the DOCX file.
    """
    doc = Document(docx_file)
    full_text = []

    for element in doc.element.body:
        # Para los párrafos
        if element.tag.endswith('p'):
            full_text.append(' '.join([node.text for node in element.iter() if node.text]))

        # Para las tablas
        elif element.tag.endswith('tbl'):
            table_texts = []
            for row in element.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                if row.text:
                    table_texts.append(row.text)
            full_text.append(' | '.join(table_texts))  # Separa el contenido de las celdas con ' | '

    return ' '.join(full_text)




# Definir la ruta donde se guardarán los archivos cargados
UPLOAD_FOLDER = '/save'
# Si deseas especificar la ruta de carga en la misma carpeta donde se encuentra este script, puedes hacerlo así:
# UPLOAD_FOLDER = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'save')
CORS(app)  # Habilita CORS para toda la aplicación


# Las funciones extract_text_from_pdf, extract_text_from_txt, extract_text_from_doc y extract_text_from_docx se mantienen igual.

@app.route("/api/appPy")
def index():
  #return render_template("index.html")
  return jsonify({"message": "El servidor Python está activo y listo para recibir solicitudes."})


@app.route("/extract_text", methods=["POST"])
def extract_text():
  if 'file' in request.files:
    file = request.files["file"]
    print(type(file))
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    if file.content_type == "application/pdf":
      text = extract_text_from_pdf(filepath)
    elif file.content_type == "text/plain":
      text = extract_text_from_txt(filepath)
    elif file.content_type == "application/msword":
      text = extract_text_from_doc(filepath)
    elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
      text = extract_text_from_docx(filepath)

    elif filename.endswith('.odt'):
      text = extract_text_from_odt(filepath)
    else:
      #return "Unsupported file type."
      return jsonify({"error": "Unsupported file type."})

    os.remove(filepath)
    print(text)
    return jsonify({"text": text})


  else:
    return "No file part in the request."

if __name__ == "__main__":
  app.run(debug=True)
